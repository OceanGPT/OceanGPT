import streamlit as st
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_huggingface.embeddings import HuggingFaceEmbeddings
from langchain_core.documents import Document
from langchain.chains.retrieval import create_retrieval_chain
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.messages import HumanMessage, SystemMessage,BaseMessage,AIMessage
import docx
from PyPDF2 import PdfReader
from transformers import AutoTokenizer, AutoModelForCausalLM, BitsAndBytesConfig
import torch
from langchain_core.language_models import BaseChatModel
from langchain_core.outputs import ChatResult, ChatGeneration
from langchain_core.callbacks import CallbackManagerForLLMRun
from typing import Any, Dict, List, Optional, Type, Iterator
import re

# 添加自定义模型类（放在其他imports下面）
class CustomChatModel(BaseChatModel):
    tokenizer: Any = None
    model: Any = None
    
    def __init__(self, tokenizer: Any, model: Any, **kwargs):
        super().__init__(tokenizer=tokenizer, model=model,**kwargs)
        self.tokenizer = tokenizer
        self.model = model

    def _generate(
        self,
        messages: List[BaseMessage],  # 修改1: 使用BaseMessage类型
        stop: Optional[List[str]] = None,
        run_manager: Optional[CallbackManagerForLLMRun] = None,
        **kwargs
    ) -> ChatResult:
        # 转换消息格式
        formatted_messages = []
        for msg in messages:
            # 修改2: 正确访问消息属性
            if isinstance(msg, SystemMessage):
                role = "system"
            elif isinstance(msg, HumanMessage):
                role = "user"
            elif isinstance(msg, AIMessage):
                role = "assistant"
            else:
                role = "user"  # 默认处理
            formatted_messages.append({
                "role": role,
                "content": msg.content  # 直接访问content属性
            })
        
        # 使用原有的模型调用逻辑
        text = self.tokenizer.apply_chat_template(
            formatted_messages,
            tokenize=False,
            add_generation_prompt=True
        )
        
        model_inputs = self.tokenizer([text], return_tensors="pt").to(self.model.device)
        generated_ids = self.model.generate(
            **model_inputs,
            max_new_tokens=kwargs.get("max_new_tokens", 8192)
        )
        
        output_ids = generated_ids[0][len(model_inputs.input_ids[0]):].tolist()
        try:
            index = len(output_ids) - output_ids[::-1].index(151668)
        except ValueError:
            index = 0
        content = self.tokenizer.decode(output_ids[index:], skip_special_tokens=True).strip("\n")
        
        # 封装为LangChain格式
        message = AIMessage(content=content)
        generation = ChatGeneration(message=message)
        return ChatResult(generations=[generation])

    @property
    def _llm_type(self) -> str:
        return "ocean-gpt-custom"

# 配置API密钥（如果使用OpenAI）
# os.environ['OPENAI_API_KEY']='xxx'
# os.environ['OPENAI_BASE_URL']='xxx' # 看你的情况

st.set_page_config(page_title="Chat with Documents", page_icon=":robot:", layout="wide")

st.markdown(
    """<style>
.chat-message {
    padding: 1.5rem; border-radius: 0.5rem; margin-bottom: 2rem; display: flex
}
.chat-message.user {
    background-color: #2b313e
}
.chat-message.bot {
    background-color: #475063
}
.chat-message .avatar {
  width: 20%;
}
.chat-message .avatar img {
  max-width: 78px;
  max-height: 78px;
}
.chat-message .message {
  width: 80%;
  padding: 0 1.5rem;
  color: #fff;
}

.stDeployButton {
            visibility: hidden;
        }
#MainMenu {visibility: hidden;}
footer {visibility: hidden;}

.block-container {
    padding: 2rem 4rem 2rem 4rem;
}

.st-emotion-cache-16txtl3 {
    padding: 3rem 1.5rem;
}
</style>
# """,
    unsafe_allow_html=True,
)

bot_template = """
<div class="chat-message bot">
    <div class="avatar">
        <img src="http://oceangpt.org/wp-content/uploads/2025/04/logo-700.png" style="max-height: 150px; max-width: 150px;">
    </div>
    <div class="message">{{MSG}}</div>
</div>
"""

user_template = """
<div class="chat-message user">
    <div class="avatar">
        <img src="https://www.shareicon.net/data/512x512/2015/09/18/103160_man_512x512.png" >
    </div>    
    <div class="message">{{MSG}}</div>
</div>
"""

def get_pdf_text(pdf_docs):
    docs = []
    text = ""
    for document in pdf_docs:
        if document.type == "application/pdf":
            pdf_reader = PdfReader(document)
            for idx, page in enumerate(pdf_reader.pages):
                temp = page.extract_text()
                clean_text = re.sub(r"^—\d+—", "", temp)
                # 为PDF页面创建Document对象，包含来源信息
                docs.append(
                    Document(
                        page_content=clean_text,
                        metadata={
                            "source": f"{document.name} - 第{idx+1}页"
                        },
                    )
                )
        elif (
            document.type
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ):
            doc = docx.Document(document)
            for idx, paragraph in enumerate(doc.paragraphs):
                docs.append(
                    Document(
                        page_content=paragraph.text,
                        metadata={
                            "source": f"{document.name} in paragraph {idx}"
                        },
                    )
                )
        elif document.type == "text/plain":
            text = document.getvalue().decode("utf-8")
            docs.append(Document(page_content=text, metadata={
                "source": document.name
            }))
    
    st.info("文档处理完成")
    return docs


def get_text_chunks(docs):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=512, chunk_overlap=64)
  
    docs_chunks = text_splitter.split_documents(docs)
    
    return docs_chunks


def get_vectorstore(docs_chunks):
    """获取向量存储，使用本地嵌入模型"""
    try:
        embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-mpnet-base-v2")
        
        # 如果输入是Document对象列表，直接使用
        vectorstore = FAISS.from_documents(docs_chunks, embedding=embeddings)
        return vectorstore
    except Exception as e:
        st.error(f"本地嵌入 模型加载失败: {str(e)}")
        return None


def clean_response(answer):
    """清理回答中的提示词和文档内容，直接提取Helpful Answer:后的内容"""
    # 查找 "Helpful Answer:" 并提取其后的内容
    if "Helpful Answer:" in answer:
        # 分割并获取Helpful Answer:后的内容
        parts = answer.split("Helpful Answer:")
        if len(parts) > 1:
            helpful_answer = parts[1].strip()
            # 移除可能的引号和多余字符
            helpful_answer = helpful_answer.strip("'\"")
            return helpful_answer
    
    # 如果没有找到Helpful Answer:，尝试其他模式
    if "AI:" in answer:
        parts = answer.split("AI:")
        if len(parts) > 1:
            return parts[1].strip().strip("'\"")
    
    # 如果都没有找到，返回原始内容
    return answer.strip()


def show_history():
    chat_history = st.session_state.get("chat_history", [])
    for i, message in enumerate(chat_history):
        role, content = message
        if role == "user":
            st.write(user_template.replace("{{MSG}}", content), unsafe_allow_html=True)
        else:
            st.write(bot_template.replace("{{MSG}}", content), unsafe_allow_html=True)

    # 添加滚动锚点
    st.markdown("<div id='bottom-history'></div>", unsafe_allow_html=True)
    st.markdown("""
    <script>
        var elem = document.getElementById("bottom-history");
        if (elem) elem.scrollIntoView({behavior: "smooth"});
    </script>
    """, unsafe_allow_html=True)



# 修改后的模型加载函数
@st.cache_resource(show_spinner="正在加载大模型...")
def load_cached_model(model_path: str):
    bnb_config = BitsAndBytesConfig(
    load_in_4bit=True,               # int4
    bnb_4bit_use_double_quant=True,
    bnb_4bit_quant_type="nf4",       # nf4精度更高
    bnb_4bit_compute_dtype="float16" # 推理用 float16
    )
    tokenizer = AutoTokenizer.from_pretrained(model_path)
    model = AutoModelForCausalLM.from_pretrained(
        model_path,
        quantization_config=bnb_config,
        torch_dtype="auto",
        device_map="auto"
    )
    return tokenizer, model

# 修改后的handle_userinput_auto函数
def handle_userinput_auto(user_question):
    if st.session_state.vectorstore is not None:
        # RAG模式 - 使用自定义模型
        prompt = ChatPromptTemplate.from_template("""根据以下给定的参考资料回答问题。如果参考资料中没有相关信息，请回答“我无法从资料中找到答案”，不要回答参考资料中未提及的信息。请用简短的一段话来回答给定的问题。
        参考资料：
        <context>
        {context}
        </context>
        问题：{input}""")
        
        # 使用自定义模型创建文档链
        document_chain = create_stuff_documents_chain(
            st.session_state.langchain_model,
            prompt
        )
        
        vector = st.session_state.vectorstore
        retriever = vector.as_retriever(search_kwargs={"k": 3})  # 检索3个最相关的文档片段
        retrieval_chain = create_retrieval_chain(retriever, document_chain)
        
        # 调用RAG链
        response = retrieval_chain.invoke({"input": user_question})
        cleaned_answer = response["answer"]
        
        # 获取检索到的文档片段
        retrieved_docs = response.get("context", [])
        
        # 构建包含来源的答案
        final_answer = cleaned_answer
        
        # 如果有检索到的文档，添加来源信息
        if retrieved_docs:
            sources_info = "\n\n**参考来源：**\n"
            for i, doc in enumerate(retrieved_docs, 1):
                # 截取文档内容的前100个字符作为预览
                content_preview = doc.page_content[:100] + "..." if len(doc.page_content) > 100 else doc.page_content
                source = doc.metadata.get("source", "未知来源")
                sources_info += f"{i}. **{source}**\n   {content_preview}\n\n"
            
            final_answer += sources_info
    else:
        # 普通对话模式 - 使用原始方法
        tokenizer = st.session_state.tokenizer
        model = st.session_state.llm
        messages = [{"role": "user", "content": user_question}]
        text = tokenizer.apply_chat_template(
            messages,
            tokenize=False,
            add_generation_prompt=True,
            enable_thinking=False
        )
        model_inputs = tokenizer([text], return_tensors="pt").to(model.device)
        generated_ids = model.generate(
            **model_inputs,
            max_new_tokens=8192
        )
        output_ids = generated_ids[0][len(model_inputs.input_ids[0]):].tolist()
        try:
            index = len(output_ids) - output_ids[::-1].index(151668)
        except ValueError:
            index = 0
        content = tokenizer.decode(output_ids[index:], skip_special_tokens=True).strip("\n")
        final_answer = clean_response(content)
    
    # 更新对话历史和UI保持不变
    st.session_state.chat_history.append(("user", user_question))
    st.session_state.chat_history.append(("assistant", final_answer))
    # st.write(user_template.replace("{{MSG}}", user_question), unsafe_allow_html=True)
    # st.write(bot_template.replace("{{MSG}}", final_answer), unsafe_allow_html=True)

def main():
    model_path = "/home/gaosong/OceanGPT-lora"
    tokenizer, model = load_cached_model(model_path)
    # 将模型放入 session_state
    if "llm" not in st.session_state:
        st.session_state.llm = model
        st.session_state.tokenizer = tokenizer
        st.session_state.llm_path = model_path
        st.session_state.langchain_model = CustomChatModel(
            tokenizer=tokenizer,
            model=model
        )
    st.header("通过定制化 OceanGPT·沧渊，对文档内容进行精问答")
    with st.sidebar:
        if "conversation" not in st.session_state:
            st.session_state.conversation = None
        if "chat_history" not in st.session_state:
            st.session_state.chat_history = []
        if "vectorstore" not in st.session_state:
            st.session_state.vectorstore = None
        if "last_uploaded_files" not in st.session_state:
            st.session_state.last_uploaded_files = None

        st.title("文档管理")
        pdf_docs = st.file_uploader(
            "请拖拽文件到下方区域，或点击按钮选择文件（支持PDF、TXT、DOC、DOCX，单文件最大200MB）",
            type=["pdf", "txt", "doc", "docx"],
            accept_multiple_files=True,
        )

        # 只在文件变化时处理
        current_files = tuple([f.name for f in pdf_docs]) if pdf_docs else None
        if current_files != st.session_state.last_uploaded_files:
            if pdf_docs:
                with st.spinner("正在处理文档..."):
                    docs = get_pdf_text(pdf_docs)  # 现在返回Document对象列表
                    docs_chunks = get_text_chunks(docs)  # 处理Document对象
                    vectorstore = get_vectorstore(docs_chunks)
                    if vectorstore is not None:
                        st.session_state.vectorstore = vectorstore
                    else:
                        st.error("向量存储创建失败")
            else:
                st.session_state.vectorstore = None
            st.session_state.last_uploaded_files = current_files

        def clear_history():
            st.session_state.chat_history = []
        if st.session_state.chat_history:
            st.button("清空对话", on_click=clear_history, use_container_width=True)

    # 创建主内容容器
    with st.container():
        # 展示历史记录的区域（位于页面顶部）
        with st.container(height=600, border=True):  # 调整高度为500
            show_history()
    
    # 创建输入容器（位于页面底部）
    with st.container():
        user_question = st.chat_input("输入点什么~", key="input_chat")
        if user_question:
            handle_userinput_auto(user_question)
            st.rerun()


if __name__ == "__main__":
    main()